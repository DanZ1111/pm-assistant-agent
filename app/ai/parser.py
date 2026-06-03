import os
import json
import base64
import re
import shutil
import subprocess
import tempfile
import time
from openai import OpenAI
from app.ai.prompts import (
    EXTRACTION_SYSTEM_PROMPT,
    VISION_EXTRACTION_SYSTEM_PROMPT,
    DUAL_MODE_INTAKE_PROMPT,
    JOURNAL_SUMMARY_PROMPT,
    BUSINESS_PLAN_EXTRACT_PROMPT,
)

_client: OpenAI | None = None

NON_USD_CURRENCY_RE = re.compile(r"(?:\bRMB\b|\bCNY\b|人民币|元|¥)", re.IGNORECASE)
USD_CURRENCY_RE = re.compile(r"(?:\$|\bUSD\b|\bUS\$|美元)", re.IGNORECASE)
RANGE_RE = re.compile(
    r"(?:\$|USD|US\$|美元)?\s*\d+(?:\.\d+)?\s*(?:-|–|—|~|～|至|到|－|\bto\b)\s*(?:\$|USD|US\$|美元)?\s*\d+(?:\.\d+)?",
    re.IGNORECASE,
)
COST_CONTEXT_RE = re.compile(r"(?:cost|factory|ex[- ]?factory|material|材料|成本|出厂)", re.IGNORECASE)
MSRP_CONTEXT_RE = re.compile(r"(?:msrp|retail|price|售价|零售|价格)", re.IGNORECASE)
PRICE_FIELDS = ("target_factory_cost", "target_msrp")


def _get_client() -> OpenAI:
    global _client
    if _client is None:
        from dotenv import load_dotenv
        load_dotenv()
        _client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"), timeout=120.0)
    return _client


def extract_project_fields(raw_text: str) -> dict:
    """Call GPT-4o with JSON mode to extract project fields from raw text.
    Returns a dict of extracted fields, or {"_error": reason} on failure.
    """
    try:
        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": EXTRACTION_SYSTEM_PROMPT},
                {"role": "user", "content": raw_text},
            ],
            max_completion_tokens=1000,
            temperature=0.1,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        return sanitize_price_fields(parsed, raw_text, source_kind="text")
    except Exception as e:
        return {"_error": str(e)}


def sanitize_price_fields(extracted: dict, source_text: str = "", source_kind: str = "text") -> dict:
    """Normalize PM-facing price fields as text.

    Project-level target cost/MSRP are planning expressions, not pure math
    inputs. Preserve ranges, currencies, and qualifiers when the source has
    them; fall back to stringifying the model value.
    """
    if not isinstance(extracted, dict):
        return extracted

    cleaned = dict(extracted)
    cleaned.pop("_warnings", None)

    for field in PRICE_FIELDS:
        value = cleaned.get(field)
        expression = None
        if source_kind != "image":
            expression = _source_expression_for_field(source_text, field, value)
        if expression:
            cleaned[field] = expression
        elif value not in (None, ""):
            cleaned[field] = str(value).strip()
        else:
            cleaned.pop(field, None)
    return cleaned


def _number_variants(value) -> set[str]:
    variants = {str(value).strip()}
    try:
        numeric = float(str(value).replace(",", ""))
    except (TypeError, ValueError):
        return {v for v in variants if v}
    variants.add(f"{numeric:g}")
    variants.add(f"{numeric:.1f}")
    variants.add(f"{numeric:.2f}")
    if numeric.is_integer():
        variants.add(str(int(numeric)))
    return {v for v in variants if v}


def _windows_around_value(source_text: str, value, radius: int = 48) -> list[str]:
    source = source_text or ""
    windows = []
    for variant in _number_variants(value):
        pattern = re.compile(rf"(?<![\d.]){re.escape(variant)}(?![\d.])")
        for match in pattern.finditer(source):
            start = max(0, match.start() - radius)
            end = min(len(source), match.end() + radius)
            windows.append(source[start:end])
    return windows


def _source_expression_for_field(source_text: str, field: str, value=None) -> str | None:
    line_expression = _line_expression_for_field(source_text, field)
    if line_expression:
        return line_expression

    if value not in (None, ""):
        for window in _windows_around_value(source_text, value, radius=32):
            range_match = RANGE_RE.search(window)
            if range_match:
                return _clean_price_expression(range_match.group(0))
            amount_match = re.search(
                r"(?:under|around|about|approx\.?|approximately|约|大约|低于|小于)?\s*"
                r"(?:\$|USD|US\$|RMB|CNY|人民币|美元|元|¥)?\s*"
                r"\d+(?:,\d{3})*(?:\.\d+)?\s*"
                r"(?:\$|USD|US\$|RMB|CNY|人民币|美元|元|¥)?"
                r"(?:\s*出厂)?",
                window,
                re.IGNORECASE,
            )
            if amount_match:
                return _clean_price_expression(amount_match.group(0))

    return None


def _line_expression_for_field(source_text: str, field: str) -> str | None:
    source = source_text or ""
    context_re = COST_CONTEXT_RE if field == "target_factory_cost" else MSRP_CONTEXT_RE
    for raw_line in source.splitlines():
        line = raw_line.strip(" \t-*•")
        if not line or not context_re.search(line):
            continue
        if not re.search(r"\d", line):
            continue
        segment = line
        if field == "target_msrp":
            match = re.search(r"(?:msrp|retail|售价|零售|价格|区间)", segment, re.IGNORECASE)
            if match:
                segment = segment[match.end():]
        else:
            match = re.search(r"(?:target\s+factory\s+cost|factory\s+cost|ex[- ]?factory|material\s+cost|材料成本|成本|出厂)", segment, re.IGNORECASE)
            if match:
                segment = segment[match.end():]
        if "：" in segment:
            segment = segment.split("：", 1)[1].strip()
        elif ":" in segment:
            segment = segment.split(":", 1)[1].strip()
        if field == "target_factory_cost":
            segment = re.split(r"\bMSRP\b|retail|售价|零售|价格", segment, maxsplit=1, flags=re.IGNORECASE)[0]
        if field == "target_msrp":
            segment = re.split(r"factory|材料|成本|出厂", segment, maxsplit=1, flags=re.IGNORECASE)[0]
        range_match = RANGE_RE.search(segment)
        if range_match:
            return _clean_price_expression(range_match.group(0))
        amount_match = re.search(
            r"(?:under|around|about|approx\.?|approximately|约|大约|低于|小于)?\s*"
            r"(?:\$|USD|US\$|RMB|CNY|人民币|美元|元|¥)?\s*"
            r"\d+(?:,\d{3})*(?:\.\d+)?\s*"
            r"(?:\$|USD|US\$|RMB|CNY|人民币|美元|元|¥)?"
            r"(?:\s*出厂)?",
            segment,
            re.IGNORECASE,
        )
        if amount_match:
            return _clean_price_expression(amount_match.group(0))
        if context_re.search(segment) and len(segment) > 48:
            continue
        return _clean_price_expression(segment)
    return None


def _clean_price_expression(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip(" \t:：；;,.，。")).strip()


def _extract_pdf_text(file_path: str, max_pages: int = 10) -> str:
    """Pure text extraction from a PDF — no AI call.
    Returns the joined text (may be empty if scanned)."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    pages = reader.pages[:max_pages]
    return "\n\n".join(page.extract_text() or "" for page in pages).strip()


def _extract_docx_text(file_path: str) -> str:
    """Pure text extraction from a .docx — paragraphs + table cells."""
    from docx import Document
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    return "\n\n".join(parts).strip()


def _extract_doc_text(file_path: str) -> str | None:
    """Convert legacy .doc → .docx using LibreOffice (soffice), then extract.
    Returns None if soffice isn't available or conversion fails — caller
    treats None as 'unsupported, surface friendly error to user.'"""
    if not shutil.which("soffice"):
        return None
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx",
                 file_path, "--outdir", tmpdir],
                capture_output=True, timeout=60,
            )
            if result.returncode != 0:
                return None
            base = os.path.splitext(os.path.basename(file_path))[0]
            converted = os.path.join(tmpdir, f"{base}.docx")
            if not os.path.exists(converted):
                return None
            return _extract_docx_text(converted)
    except Exception:
        return None


def extract_from_pdf(file_path: str) -> dict:
    """Extract text from a PDF and run field extraction.
    Returns {"extracted": dict, "raw_text": str} or {"_error": str}.
    """
    try:
        raw_text = _extract_pdf_text(file_path)
        if not raw_text:
            return {"_error": "Could not extract text from PDF — it may be a scanned image or empty."}
        extracted = extract_project_fields(raw_text)
        if "_error" in extracted:
            return extracted
        return {"extracted": extracted, "raw_text": raw_text}
    except Exception as e:
        return {"_error": str(e)}


def extract_from_image(file_path: str) -> dict:
    """Use vision API to extract fields and generate a summary from an image.
    Returns {"extracted": dict, "ai_summary": str} or {"_error": str}.
    """
    try:
        ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else "png"
        mime = f"image/{ext}" if ext in ("png", "jpg", "jpeg", "gif", "webp") else "image/png"
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": VISION_EXTRACTION_SYSTEM_PROMPT},
                {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": "Analyze this product image and extract information."},
                ]},
            ],
            max_completion_tokens=1000,
            temperature=0.1,
        )
        raw = response.choices[0].message.content or "{}"
        result = json.loads(raw)
        result = sanitize_price_fields(result, "", source_kind="image")
        ai_summary = result.pop("ai_summary", "")
        return {"extracted": result, "ai_summary": ai_summary}
    except Exception as e:
        return {"_error": str(e)}


def answer_help_question(question: str) -> str:
    """Answer a question about how to use PM Product Tracker using USER_GUIDE.md and CHANGELOG.md."""
    root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

    def _read(name, limit=8000):
        path = os.path.join(root, name)
        try:
            with open(path, "r", encoding="utf-8") as f:
                return f.read()[:limit]
        except Exception:
            return f"({name} not found)"

    user_guide = _read("USER_GUIDE.md")
    changelog = _read("CHANGELOG.md")

    system_prompt = (
        "You are a helpful onboarding assistant for PM Product Tracker — "
        "an internal product management tool for a knife and product development company.\n\n"
        "Answer the user's question based only on the documentation below. "
        "Be concise (under 200 words). Use numbered steps for how-to questions. "
        "Use bullet points for lists. Never invent features not in the docs. "
        "If the question isn't covered, say so and suggest they check with the team.\n\n"
        f"---\nUSER GUIDE:\n{user_guide}\n\n---\nCHANGELOG:\n{changelog}"
    )

    try:
        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question},
            ],
            max_completion_tokens=600,
            temperature=0.2,
        )
        return response.choices[0].message.content or "I couldn't generate an answer."
    except Exception as e:
        return f"Sorry, I couldn't answer that right now: {e}"


def extract_intake(raw_text: str) -> dict:
    """Dual-mode classifier: returns {classification: 'project'|'idea',
    project_fields: {...}, idea_fields: {...}}. Defaults to 'idea' on
    ambiguous text per product policy.
    """
    try:
        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": DUAL_MODE_INTAKE_PROMPT},
                {"role": "user", "content": raw_text},
            ],
            max_completion_tokens=1200,
            temperature=0.1,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        # Normalize the shape
        cls = parsed.get("classification")
        if cls not in ("project", "idea"):
            cls = "idea"  # default to idea on missing/invalid classification
        project_fields = parsed.get("project_fields") or {}
        if cls == "project":
            project_fields = sanitize_price_fields(project_fields, raw_text, source_kind="text")
        return {
            "classification": cls,
            "project_fields": project_fields,
            "idea_fields": parsed.get("idea_fields") or {},
        }
    except Exception as e:
        return {"_error": str(e)}


def extract_thesis_and_inspirations(file_path: str, file_type: str) -> dict:
    """Build 15 — Extract a product thesis + inspirations from an uploaded
    business plan. One-time AI call; the result is persisted as an
    AIMessage row so the preview page can re-render without re-running AI.

    file_type: 'pdf' | 'docx' | 'doc' | 'image'
    Returns:
      {
        "thesis": str,
        "inspirations": [{name, description, idea_type, source, source_detail}],
        "raw_text_preview": str,
        "duration_seconds": float,
        "model": str,
      }
      or {"_error": str} on any failure.
    """
    started = time.time()
    file_type = (file_type or "").lower().strip()

    try:
        if file_type == "pdf":
            raw_text = _extract_pdf_text(file_path)
            if not raw_text:
                return {"_error": "Could not extract text from PDF — it may be a scanned image or empty."}
        elif file_type == "docx":
            raw_text = _extract_docx_text(file_path)
            if not raw_text:
                return {"_error": "DOCX file appears to be empty."}
        elif file_type == "doc":
            raw_text = _extract_doc_text(file_path)
            if raw_text is None:
                return {"_error": "DOC format requires LibreOffice on the server. Please save as DOCX or PDF and re-upload."}
            if not raw_text:
                return {"_error": "DOC file appears to be empty after conversion."}
        elif file_type == "image":
            return _extract_thesis_from_image(file_path, started)
        else:
            return {"_error": f"Unsupported file type '{file_type}'. Use PDF, DOCX, DOC, or image."}

        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": BUSINESS_PLAN_EXTRACT_PROMPT},
                {"role": "user", "content": raw_text[:30000]},
            ],
            max_completion_tokens=1500,
            temperature=0.2,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        thesis = (parsed.get("thesis") or "").strip()
        inspirations = parsed.get("inspirations") or []
        if not isinstance(inspirations, list):
            inspirations = []
        return {
            "thesis": thesis,
            "inspirations": inspirations,
            "raw_text_preview": raw_text[:600],
            "duration_seconds": round(time.time() - started, 2),
            "model": "gpt-5.4",
        }
    except Exception as e:
        return {"_error": str(e)}


def _extract_thesis_from_image(file_path: str, started: float) -> dict:
    """Vision-based thesis + inspirations extraction from an image."""
    try:
        ext = file_path.rsplit(".", 1)[-1].lower() if "." in file_path else "png"
        mime = f"image/{ext}" if ext in ("png", "jpg", "jpeg", "gif", "webp") else "image/png"
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")

        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": BUSINESS_PLAN_EXTRACT_PROMPT},
                {"role": "user", "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    {"type": "text", "text": "This is an image of a product brief or business plan page. Extract thesis and inspirations."},
                ]},
            ],
            max_completion_tokens=1500,
            temperature=0.2,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        thesis = (parsed.get("thesis") or "").strip()
        inspirations = parsed.get("inspirations") or []
        if not isinstance(inspirations, list):
            inspirations = []
        return {
            "thesis": thesis,
            "inspirations": inspirations,
            "raw_text_preview": "(image source — no raw text)",
            "duration_seconds": round(time.time() - started, 2),
            "model": "gpt-5.4",
        }
    except Exception as e:
        return {"_error": str(e)}


def summarize_journal_entry(entry_text: str) -> dict:
    """Summarize a Project Journal entry into {'title': str, 'summary': str}.
    Returns {'_error': str} on failure so the caller can preserve the existing
    title/summary instead of overwriting with garbage.
    """
    if not entry_text or not entry_text.strip():
        return {"_error": "Entry text is empty."}
    try:
        client = _get_client()
        response = client.chat.completions.create(
            model="gpt-5.4",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": JOURNAL_SUMMARY_PROMPT},
                {"role": "user", "content": entry_text},
            ],
            max_completion_tokens=300,
            temperature=0.2,
        )
        raw = response.choices[0].message.content or "{}"
        parsed = json.loads(raw)
        title = (parsed.get("title") or "").strip()
        summary = (parsed.get("summary") or "").strip()
        if not title or not summary:
            return {"_error": "AI returned an incomplete summary."}
        return {"title": title, "summary": summary}
    except Exception as e:
        return {"_error": str(e)}
